import os, pytesseract
from PIL import Image
from docx import Document

dir_path = (r'<absolute-path-to-screenshots_folder>')
# you can remove the following line if you've passed Tesseract to PATH
pytesseract.pytesseract.tesseract_cmd = r'<absolute-path-to-tesseract-executable>'

def transform_to_text():
  count = 0
  document = Document()
  file_name = input("How do you want to name your file (without the suffix)? ")
  disallowed_characters = list(':/;?<>|\\*"')
  while True:
    file_name_chars = list(file_name)
    if any(char in disallowed_characters for char in file_name_chars):
     file_name = input("Sorry, can't have any file name containing the following characters: :/;?<>|\\*\" \nPlease choose something else: ")
    else:
      break
  document.add_heading(file_name, 0)
  line_string = "-------------------------------------------------------------------\n"
  for file in os.listdir(dir_path):
      try:
        output = pytesseract.image_to_string(Image.open(os.path.join(dir_path, file))) + "\n"
        document.add_paragraph(output)
        document.add_paragraph(line_string)
        print("Screenshot successfully transcribed!")
        count += 1
      except:
        print(f"Could not transcribe the following file: {file}")
        continue
  if count > 0:
    print(f"{count} screenshot{'s' if count > 1 else ''} transcribed!")
    document.save(f"{file_name}.docx")
  else:
     print("Sorry, could not transcribe any files!")

if __name__ == "__main__":
   transform_to_text()
